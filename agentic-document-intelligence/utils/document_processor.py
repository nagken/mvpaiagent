"""
Document Processing Utilities for Agentic Document Intelligence System
Handles PDF extraction, text processing, and metadata extraction
"""

import asyncio
import logging
import mimetypes
import os
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import hashlib
import json

# Document processing imports (will be available after pip install)
try:
    import PyPDF2
    import fitz  # PyMuPDF
    import pdfplumber
    from docx import Document
except ImportError:
    logging.warning("Document processing libraries not installed. Run: pip install -r requirements.txt")

import pandas as pd
import numpy as np
from dataclasses import dataclass
from datetime import datetime


@dataclass
class DocumentMetadata:
    """Document metadata structure"""
    filename: str
    file_path: str
    file_size: bytes
    file_hash: str
    mime_type: str
    created_at: datetime
    modified_at: datetime
    page_count: Optional[int] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    creator: Optional[str] = None
    language: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            'filename': self.filename,
            'file_path': self.file_path,
            'file_size': self.file_size,
            'file_hash': self.file_hash,
            'mime_type': self.mime_type,
            'created_at': self.created_at.isoformat(),
            'modified_at': self.modified_at.isoformat(),
            'page_count': self.page_count,
            'author': self.author,
            'title': self.title,
            'subject': self.subject,
            'creator': self.creator,
            'language': self.language
        }


@dataclass
class DocumentChunk:
    """Document text chunk with metadata"""
    chunk_id: str
    content: str
    page_number: Optional[int]
    chunk_index: int
    start_char: int
    end_char: int
    document_hash: str
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            'chunk_id': self.chunk_id,
            'content': self.content,
            'page_number': self.page_number,
            'chunk_index': self.chunk_index,
            'start_char': self.start_char,
            'end_char': self.end_char,
            'document_hash': self.document_hash,
            'metadata': self.metadata
        }


class DocumentProcessor:
    """Handles document processing and text extraction"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger("document_processor")
        self.supported_formats = ['.pdf', '.docx', '.txt']
        
        # Chunk configuration
        self.chunk_size = config.get('chunk_size', 1000)
        self.chunk_overlap = config.get('chunk_overlap', 200)
    
    async def process_document(self, file_path: str) -> Tuple[DocumentMetadata, List[DocumentChunk]]:
        """Process a document and return metadata and chunks"""
        try:
            # Extract metadata
            metadata = await self._extract_metadata(file_path)
            
            # Extract text content
            text_content = await self._extract_text(file_path, metadata.mime_type)
            
            # Create chunks
            chunks = await self._create_chunks(text_content, metadata.file_hash)
            
            self.logger.info(
                f"Processed {metadata.filename}: {len(chunks)} chunks, "
                f"{len(text_content)} characters"
            )
            
            return metadata, chunks
            
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    async def _extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract document metadata"""
        path = Path(file_path)
        stat = path.stat()
        
        # Basic file metadata
        metadata = DocumentMetadata(
            filename=path.name,
            file_path=str(path.absolute()),
            file_size=stat.st_size,
            file_hash=await self._calculate_file_hash(file_path),
            mime_type=mimetypes.guess_type(file_path)[0] or 'application/octet-stream',
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime)
        )
        
        # Extract format-specific metadata
        if metadata.mime_type == 'application/pdf':
            await self._extract_pdf_metadata(file_path, metadata)
        elif metadata.mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            await self._extract_docx_metadata(file_path, metadata)
        
        return metadata
    
    async def _extract_pdf_metadata(self, file_path: str, metadata: DocumentMetadata):
        """Extract PDF-specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata.page_count = len(pdf_reader.pages)
                
                if pdf_reader.metadata:
                    metadata.author = pdf_reader.metadata.get('/Author')
                    metadata.title = pdf_reader.metadata.get('/Title')
                    metadata.subject = pdf_reader.metadata.get('/Subject')
                    metadata.creator = pdf_reader.metadata.get('/Creator')
                    
        except Exception as e:
            self.logger.warning(f"Could not extract PDF metadata: {str(e)}")
    
    async def _extract_docx_metadata(self, file_path: str, metadata: DocumentMetadata):
        """Extract DOCX-specific metadata"""
        try:
            doc = Document(file_path)
            properties = doc.core_properties
            
            metadata.author = properties.author
            metadata.title = properties.title
            metadata.subject = properties.subject
            
        except Exception as e:
            self.logger.warning(f"Could not extract DOCX metadata: {str(e)}")
    
    async def _extract_text(self, file_path: str, mime_type: str) -> str:
        """Extract text content from document"""
        if mime_type == 'application/pdf':
            return await self._extract_pdf_text(file_path)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return await self._extract_docx_text(file_path)
        elif mime_type == 'text/plain':
            return await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods for best results"""
        text_content = ""
        
        try:
            # Method 1: PyMuPDF (best for most PDFs)
            doc = fitz.open(file_path)
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            doc.close()
            
        except Exception as e:
            self.logger.warning(f"PyMuPDF extraction failed: {str(e)}")
            
            try:
                # Method 2: pdfplumber (good for tables and complex layouts)
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text
                            
            except Exception as e2:
                self.logger.warning(f"pdfplumber extraction failed: {str(e2)}")
                
                try:
                    # Method 3: PyPDF2 (fallback)
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page in pdf_reader.pages:
                            text_content += page.extract_text()
                            
                except Exception as e3:
                    self.logger.error(f"All PDF extraction methods failed: {str(e3)}")
                    raise
        
        return text_content.strip()
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {str(e)}")
            raise
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file")
    
    async def _create_chunks(self, text: str, document_hash: str) -> List[DocumentChunk]:
        """Split text into overlapping chunks"""
        chunks = []
        
        if len(text) <= self.chunk_size:
            # Single chunk for small documents
            chunk = DocumentChunk(
                chunk_id=f"{document_hash}_chunk_0",
                content=text,
                page_number=None,
                chunk_index=0,
                start_char=0,
                end_char=len(text),
                document_hash=document_hash,
                metadata={'total_chunks': 1}
            )
            chunks.append(chunk)
            return chunks
        
        # Split into overlapping chunks
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end within overlap distance
                for i in range(end - self.chunk_overlap, end):
                    if i > start and text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk_content = text[start:end].strip()
            
            if chunk_content:  # Only add non-empty chunks
                chunk = DocumentChunk(
                    chunk_id=f"{document_hash}_chunk_{chunk_index}",
                    content=chunk_content,
                    page_number=None,  # Could be enhanced to track pages
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end,
                    document_hash=document_hash,
                    metadata={'chunk_length': len(chunk_content)}
                )
                chunks.append(chunk)
                chunk_index += 1
            
            # Move start position (with overlap)
            start = end - self.chunk_overlap if end < len(text) else end
        
        # Update total chunk count in metadata
        for chunk in chunks:
            chunk.metadata['total_chunks'] = len(chunks)
        
        return chunks
    
    async def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        
        with open(file_path, 'rb') as file:
            for chunk in iter(lambda: file.read(4096), b""):
                hash_sha256.update(chunk)
        
        return hash_sha256.hexdigest()
    
    async def batch_process_documents(self, file_paths: List[str]) -> Dict[str, Tuple[DocumentMetadata, List[DocumentChunk]]]:
        """Process multiple documents in batch"""
        results = {}
        
        # Process documents concurrently
        semaphore = asyncio.Semaphore(self.config.get('max_concurrent', 3))
        
        async def process_single(file_path: str):
            async with semaphore:
                try:
                    metadata, chunks = await self.process_document(file_path)
                    results[file_path] = (metadata, chunks)
                except Exception as e:
                    self.logger.error(f"Failed to process {file_path}: {str(e)}")
                    results[file_path] = None
        
        # Execute all processing tasks
        tasks = [process_single(fp) for fp in file_paths]
        await asyncio.gather(*tasks, return_exceptions=True)
        
        successful = len([r for r in results.values() if r is not None])
        self.logger.info(f"Batch processed {successful}/{len(file_paths)} documents")
        
        return results
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported"""
        return Path(file_path).suffix.lower() in self.supported_formats
    
    async def validate_document(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """Validate document before processing"""
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                return False, "File does not exist"
            
            # Check file size
            max_size = self.config.get('max_file_size_mb', 100) * 1024 * 1024
            if path.stat().st_size > max_size:
                return False, f"File size exceeds {max_size // (1024*1024)}MB limit"
            
            # Check file format
            if not self.is_supported_format(file_path):
                return False, f"Unsupported file format: {path.suffix}"
            
            return True, None
            
        except Exception as e:
            return False, f"Validation error: {str(e)}"


class TextAnalyzer:
    """Advanced text analysis utilities"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger("text_analyzer")
    
    async def extract_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract named entities from text (simplified version)"""
        # This would integrate with spaCy or similar NLP library
        # For now, return placeholder structure
        entities = {
            'PERSON': [],
            'ORGANIZATION': [],
            'DATE': [],
            'LOCATION': [],
            'CONCEPT': []
        }
        
        # TODO: Implement actual NLP entity extraction
        # import spacy
        # nlp = spacy.load("en_core_web_sm")
        # doc = nlp(text)
        # for ent in doc.ents:
        #     if ent.label_ in entities:
        #         entities[ent.label_].append(ent.text)
        
        return entities
    
    async def analyze_sentiment(self, text: str) -> Dict[str, float]:
        """Analyze text sentiment"""
        # Placeholder for sentiment analysis
        return {
            'positive': 0.5,
            'negative': 0.2,
            'neutral': 0.3,
            'compound': 0.3
        }
    
    async def extract_keywords(self, text: str, top_k: int = 10) -> List[Tuple[str, float]]:
        """Extract key terms from text"""
        # Placeholder for keyword extraction (could use TF-IDF, RAKE, etc.)
        words = text.lower().split()
        word_freq = {}
        
        for word in words:
            if len(word) > 3:  # Filter short words
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Sort by frequency and return top_k
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return sorted_words[:top_k]
    
    async def calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate text similarity score"""
        # Simple Jaccard similarity for now
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        intersection = len(words1 & words2)
        union = len(words1 | words2)
        
        return intersection / union if union > 0 else 0.0